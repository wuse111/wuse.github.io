# coding: utf-8
import requests
from lxml import etree
from time import sleep
from random import randint



class N_spider():

    def __init__(self):
        self.user_agent = [
            "Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10_6_8; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50",
            "Mozilla/5.0 (Windows; U; Windows NT 6.1; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50",
            "Mozilla/5.0 (Windows NT 10.0; WOW64; rv:38.0) Gecko/20100101 Firefox/38.0",
            "Mozilla/5.0 (Windows NT 10.0; WOW64; Trident/7.0; .NET4.0C; .NET4.0E; .NET CLR 2.0.50727; .NET CLR 3.0.30729; .NET CLR 3.5.30729; InfoPath.3; rv:11.0) like Gecko",
            "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Trident/5.0)",
            "Mozilla/4.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)",
            "Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1)",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10.6; rv:2.0.1) Gecko/20100101 Firefox/4.0.1",
            "Mozilla/5.0 (Windows NT 6.1; rv:2.0.1) Gecko/20100101 Firefox/4.0.1",
            "Opera/9.80 (Macintosh; Intel Mac OS X 10.6.8; U; en) Presto/2.8.131 Version/11.11",
            "Opera/9.80 (Windows NT 6.1; U; en) Presto/2.8.131 Version/11.11",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_0) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Maxthon 2.0)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; TencentTraveler 4.0)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; The World)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Trident/4.0; SE 2.X MetaSr 1.0; SE 2.X MetaSr 1.0; .NET CLR 2.0.50727; SE 2.X MetaSr 1.0)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; 360SE)",
            "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Avant Browser)"]
        self.header = {
            "User-Agent": self.user_agent[randint(0, 19)]
        }
        self.front = 'https://news.swufe.edu.cn/'

    def get_texturls(self):
        page = 0
        new_list = []
        for i in range(1):
            sleep(randint(1, 4))
            if page % 5 == 0:
                self.header["User-Agent"] = self.user_agent[randint(0, 19)]
            if page == 0:
                url = "https://news.swufe.edu.cn/zhxw.htm"
            else:
                url = f"https://news.swufe.edu.cn/zhxw/{35 - i}.htm"
            page += 1
            res = requests.get(url, headers=self.header)
            if res.status_code == 200:
                response = res.content.decode()
                selector = etree.HTML(response)
                url_list = selector.xpath('//a[@class="jiequ"]/@href')
                new_list.extend(url_list)
                print(url_list)
            else:
                self.header["User-Agent"] = self.user_agent[randint(0, 19)]
                print(f'error:网页访问出错')
        return new_list

    def get_text(self, list1):
        all_text = []
        count = 0
        for i in list1:
            sleep(randint(1, 4))
            if count % 5 == 0:
                self.header["User-Agent"] = self.user_agent[randint(0, 19)]
            if '..' in i:
                complete_url = i.repacle('..', self.front)
            else:
                complete_url = self.front + i
            res = requests.get(complete_url, headers=self.header)
            count += 1
            if res.status_code == 200:
                response = res.content.decode()
                selector = etree.HTML(response)
                text_list = selector.xpath('//div[@id="vsb_content_501"]/p/text()')
                new_text = self.deal_text(text_list)
                print(new_text)
                all_text.append(new_text)
            else:
                print(f'error:网页访问出错')
        return all_text

    def deal_text(self, list1):
        new = ""
        for i in list1:
            if i == '\r\n ':
                continue
            else:
                new = new + i
        return new

    def save_word(self, list1):
        from docx import Document
        from docx.oxml.ns import qn
        document = Document()
        p = document.add_paragraph('')
        all_text = ""
        for i in list1:
            all_text = all_text + i + '\r\n'
        p = p.add_run(all_text)
        p.font.name = 'Times New Roman'
        p.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        # document.add_page_break()
        document.save('文本.docx')

    def word_cloud(self, list1):
        '''读入外面的停用词'''
        import jieba
        import wordcloud
        string = ""
        for i in list1:
            txt_list = jieba.lcut(i)
            string1 = " ".join(txt_list)
            string = string + string1
        s_list = open('stopwords.txt', 'r', encoding='utf-8').read().split('\n')
        stopwords = set(s_list)
        w = wordcloud.WordCloud(width=1300, height=1000, background_color='white', font_path='msyh.ttc',
                                stopwords=stopwords)
        w.generate(string)
        w.to_file("词云图.png")
if __name__ == '__main__':
    spider = N_spider()
    back_list = spider.get_texturls()
    ok_list = spider.get_text(back_list)
    spider.save_word(ok_list)
    spider.word_cloud(ok_list)